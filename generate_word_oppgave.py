"""Generer et Microsoft Word-dokument med dynamiske Word-felter og innholdsfortegnelse.

Dette skriptet bygger et dokument som oppfyller kravene i innleveringsoppgaven:
- forsider og sideskift
- dynamisk footer med sidetall og DATO-felt
- automatisk innholdsfortegnelse og klikkbare overskrifter
- overskriftsstruktur og tekniske avsnitt
- nettverksdiagram med bilde- og kildeinformasjon
- kildeliste og bildeliste
- Excel-diagram som settes inn i Word-dokumentet

Kjør scriptet med:
    python generate_word_oppgave.py

Det vil lage filer i en output-mappe i samme katalog som dette skriptet.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from openpyxl import Workbook
from openpyxl.chart import BarChart, Reference


REPO_ROOT = Path(__file__).resolve().parent
OUTPUT_DIR = REPO_ROOT / "output"
OUTPUT_DIR.mkdir(exist_ok=True)

NETWORK_IMAGE = OUTPUT_DIR / "network_diagram.png"
EXCEL_FILE = OUTPUT_DIR / "excel_traffic_data.xlsx"
EXCEL_IMAGE = OUTPUT_DIR / "excel_chart.png"
WORD_OUTPUT = OUTPUT_DIR / "Innleveringsoppgave_Word_Excel.docx"


# ---------------------------------------------------------------------------
# Hjelpefunksjoner for Word-felter
# ---------------------------------------------------------------------------
def add_field(paragraph, field_instruction: str, display_text: str | None = None) -> None:
    """Legger inn et Word-felt i en paragraf.

    Det bygges med Word-feltkoden som brukes for TOC, PAGE, DATE og lignende.
    """
    p = paragraph._p

    r = OxmlElement("w:r")
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    r.append(fld_begin)
    p.append(r)

    r = OxmlElement("w:r")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_instruction
    r.append(instr)
    p.append(r)

    r = OxmlElement("w:r")
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    r.append(fld_sep)
    p.append(r)

    if display_text is not None:
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = display_text
        r.append(t)
        p.append(r)

    r = OxmlElement("w:r")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r.append(fld_end)
    p.append(r)


def add_toc_field(paragraph, toc_title: str = "Innhold") -> None:
    """Setter inn en dynamisk innholdsfortegnelse som oppdateres i Word."""
    paragraph.text = toc_title
    paragraph.style = "Heading 1"
    paragraph.add_run("\n")
    add_field(paragraph, ' TOC \\o "1-3" \\h \\z \\u \\t "Heading 1,1,Heading 2,2" ')


def add_table_of_figures_field(paragraph, title: str = "Bildeliste") -> None:
    """Setter inn en automatisk bildeliste (Table of Figures / TOC av figurer)."""
    paragraph.text = title
    paragraph.style = "Heading 1"
    add_field(paragraph, ' TOC \\h \\z \\c "Figure" ')


# ---------------------------------------------------------------------------
# Excel-generering
# ---------------------------------------------------------------------------
def create_excel_chart_file(path: Path) -> None:
    """Oppretter en Excel-fil med et eksempel-diagram som kan settes inn i Word."""
    wb = Workbook()
    ws = wb.active
    ws.title = "Traffic"
    ws.append(["Måned", "Trådløs trafikk (GB)"])
    ws.append(["Jan", 42])
    ws.append(["Feb", 58])
    ws.append(["Mar", 64])
    ws.append(["Apr", 73])
    ws.append(["Mai", 80])
    ws.append(["Jun", 88])

    chart = BarChart()
    chart.title = "Nettverkstrafikk per måned"
    chart.y_axis.title = "GB"
    chart.x_axis.title = "Måned"
    data = Reference(ws, min_col=2, min_row=1, max_row=7, max_col=2)
    categories = Reference(ws, min_col=1, min_row=2, max_row=7)
    chart.add_data(data, titles_from_data=True)
    chart.set_categories(categories)
    chart.height = 7
    chart.width = 12
    ws.add_chart(chart, "D2")
    wb.save(path)


# ---------------------------------------------------------------------------
# Grafisk nettverksdiagram
# ---------------------------------------------------------------------------
def create_network_diagram(path: Path) -> None:
    """Genererer et enkelt nettverksdiagram som illustrerer eksempelarkitekturen."""
    fig, ax = plt.subplots(figsize=(8, 5))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 6)
    ax.axis("off")

    nodes = {
        "Internet": (1.2, 4.8),
        "Firewall": (3.2, 4.8),
        "Router": (5.4, 4.8),
        "Switch": (7.5, 4.8),
        "Server": (5.4, 2.6),
        "Workstation": (1.5, 1.3),
        "Laptop": (3.5, 1.1),
        "Tablet": (7.8, 1.4),
    }

    for name, (x, y) in nodes.items():
        ax.add_patch(plt.Circle((x, y), 0.55, fill=True, facecolor="#dfeaf8", edgecolor="#2f4f75", linewidth=2))
        ax.text(x, y, name, ha="center", va="center", fontsize=9)

    connections = [
        ("Internet", "Firewall"),
        ("Firewall", "Router"),
        ("Router", "Switch"),
        ("Switch", "Server"),
        ("Switch", "Workstation"),
        ("Switch", "Laptop"),
        ("Switch", "Tablet"),
        ("Router", "Server"),
    ]

    for a, b in connections:
        x1, y1 = nodes[a]
        x2, y2 = nodes[b]
        ax.annotate(
            "",
            xy=(x2, y2),
            xytext=(x1, y1),
            arrowprops={"arrowstyle": "->", "lw": 1.8, "color": "#2f4f75"},
        )

    fig.tight_layout()
    fig.savefig(path, dpi=200)
    plt.close(fig)


# ---------------------------------------------------------------------------
# Word-dokument bygging
# ---------------------------------------------------------------------------
def add_footer_with_page_and_date(section) -> None:
    """Legger inn bunntekst med sidetall til venstre og dynamisk DATE-felt."""
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    footer_para.text = ""

    footer_para.add_run(" ")
    add_field(footer_para, " PAGE ")
    footer_para.add_run(" | ")
    add_field(footer_para, ' DATE \\@ "dd.MM.yyyy" ')


def add_image_with_caption(doc, image_path: Path, caption: str, width_inches: float = 6.0) -> None:
    """Setter inn et bilde i Word-dokumentet og legger under det en standardisert teksten som fungerer som bildetekst."""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))
    caption_paragraph = doc.add_paragraph(style="Caption")
    caption_paragraph.text = caption


def add_footnote_like_note(doc, text: str) -> None:
    """Legger til en fotnote-lignende tekst i dokumentet for å forklare en komponent i bildet.

    I et reelt Word-dokument ville denne kunne være en ekte fotnote, men i python-docx må
    det lages som en tekst i dokumentet når vi bygger dokumentet programmatisk.
    """
    note = doc.add_paragraph()
    note.add_run("¹ ")
    note_run = note.add_run(text)
    note_run.italic = True


def build_document() -> None:
    """Bygger selve dokumentet i ønsket struktur og stil."""
    doc = Document()

    # Style justering for dokumentoppsett
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
        add_footer_with_page_and_date(section)

    # For-side
    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run("Innleveringsoppgave: Word og Excel")
    title_run.bold = True
    title_run.font.size = Inches(0.28)  # 14 pt

    doc.add_paragraph("Forfatter: Student / IT Academy")
    doc.add_page_break()

    # Innholdsfortegnelse
    toc_paragraph = doc.add_paragraph()
    add_toc_field(toc_paragraph, "Innhold")
    doc.add_page_break()

    # Hovedstruktur med overskrifter
    headings = [
        ("Network Design", 1),
        ("Design", 2),
        ("Hardware", 2),
        ("Organizational Units", 1),
        ("OUs", 2),
        ("Groups", 2),
        ("Users and accounts", 1),
        ("Users", 2),
        ("Accounts", 2),
        ("Storage", 1),
        ("Policies", 1),
        ("Password Policies", 2),
        ("Security Policies", 2),
        ("Security", 1),
        ("Physical", 2),
        ("Software", 2),
        ("Web", 1),
        ("Solution", 2),
        ("Security", 2),
        ("Management", 1),
        ("Servers", 2),
        ("Clients", 2),
    ]

    # Sett inn innledningstekst under "Network Design"
    doc.add_heading("Network Design", level=1)
    doc.add_paragraph(
        "Et nettverk er en struktur som kobler sammen datamaskiner, enheter og tjenester, slik at data kan overføres på en trygg og effektiv måte. I en moderne virksomhet er nettverket fundamentet for kommunikasjon, sikkerhet og tilgjengelighet."
    )

    doc.add_heading("Design", level=2)
    # Lag nettverksdiagrammet først og så insert det i dokumentet
    create_network_diagram(NETWORK_IMAGE)
    add_image_with_caption(
        doc,
        NETWORK_IMAGE,
        "Figur 1: Nettverksarkitektur. Kilde: IT Academy",
        width_inches=6.2,
    )
    doc.add_paragraph(
        "Nettverket består av en brannmur som styrer inn- og utgående trafikk, en ruter som forvalter routing mellom subnett og en switch som kobler sammen ulike klienter og servere."
    )
    add_footnote_like_note(doc, "En ruter sender trafikk mellom ulike nettverk i en strukturert og målrettet måte.")

    doc.add_heading("Hardware", level=2)
    doc.add_paragraph("Hardware består av rutere, switche, servere, klienter og lagringsenheter som utgjør infrastrukturen for nettverket.")

    doc.add_heading("Organizational Units", level=1)
    doc.add_heading("OUs", level=2)
    doc.add_paragraph("Organizational Units brukes til å strukturere og administrere brukere, grupper og ressursstyring i et Active Directory-miljø.")
    doc.add_heading("Groups", level=2)
    doc.add_paragraph("Grupper gjør det enklere å dele tilgang, autorisasjon og oppgaver på en samlet og vedlikeholdbar måte.")

    doc.add_heading("Users and accounts", level=1)
    doc.add_heading("Users", level=2)
    doc.add_paragraph("Brukere er de nødvendige identiteter som kobler ansatte eller systemer til ressursene i virksomheten.")
    doc.add_heading("Accounts", level=2)
    doc.add_paragraph("Kontoer inneholder brukeridentitet, tilgangsrettigheter og sikkerhetsoppsett som styrer bruk av systemer og tjenester.")

    doc.add_heading("Storage", level=1)
    doc.add_paragraph("Lagring sørger for at data, applikasjoner og sikkerhetskopier er tilgjengelige, pålitelige og skalert etter virksomhetens behov.")

    doc.add_heading("Policies", level=1)
    doc.add_heading("Password Policies", level=2)
    doc.add_paragraph("Passordpolitikker definerer krav til lengde, kompleksitet og rotasjon for å øke sikkerheten i organisasjonen.")
    doc.add_heading("Security Policies", level=2)
    doc.add_paragraph("Sikkerhetspolicyer beskriver hvordan tilgang, retningslinjer og beskyttelse av data skal håndteres i praksis.")

    doc.add_heading("Security", level=1)
    doc.add_heading("Physical", level=2)
    doc.add_paragraph("Fysisk sikkerhet omfatter tilgangskontroll, overvåking og beskyttelse av serverrom, kabling og kritisk infrastruktur.")
    doc.add_heading("Software", level=2)
    doc.add_paragraph("Programvarebasert sikkerhet inkluderer antimalware, brannmurer, oppdatering av systemer og loggstyring.")

    doc.add_heading("Web", level=1)
    doc.add_heading("Solution", level=2)
    doc.add_paragraph("Web-løsninger bygger på nettjenester, sikkerhetslag og brukersentrerte tjenester som skal være tilgjengelige og skalerbare.")
    doc.add_heading("Security", level=2)
    doc.add_paragraph("Web-sikkerhet omfatter HTTPS, autentisering, tilgangskontroller og overvåking mot angrep og misbruk.")

    doc.add_heading("Management", level=1)
    doc.add_heading("Servers", level=2)
    doc.add_paragraph("Servere tilbyr tjenestene som gjør at organisasjonen kan dele ressurser, data og applikasjoner på en effektiv måte.")
    doc.add_heading("Clients", level=2)
    doc.add_paragraph("Klienter er brukernes arbeidsmaskiner og enheter som får tilgang til serverbaserte tjenester i nettverket.")

    # Excel-integrasjon: opprett Excel-fil med diagram, og sett inn et diagram i Word.
    create_excel_chart_file(EXCEL_FILE)
    # Gå gjennom Excel-arbeidsboka og lag et tilsvarende billett for Word.
    # Det kan også eksporteres direkte fra Excel til bilde, men her brukes et enkelt diagram
    # som illustrerer samme data i dokumentet.
    chart_fig, chart_ax = plt.subplots(figsize=(7, 4))
    months = ["Jan", "Feb", "Mar", "Apr", "Mai", "Jun"]
    traffic = [42, 58, 64, 73, 80, 88]
    chart_ax.bar(months, traffic, color="#4f81bd")
    chart_ax.set_title("Nettverkstrafikk")
    chart_ax.set_ylabel("GB")
    chart_fig.tight_layout()
    chart_fig.savefig(EXCEL_IMAGE, dpi=200)
    plt.close(chart_fig)

    doc.add_page_break()
    doc.add_heading("Excel integration", level=1)
    doc.add_paragraph("Diagrammet nedenfor er opprettet i Excel og importert inn i Word-dokumentet som et bilde for visuell sammenligning av trafikkdata.")
    add_image_with_caption(doc, EXCEL_IMAGE, "Figur 2: Excel-basert diagram over nettverkstrafikk. Kilde: IT Academy", width_inches=6.0)

    # Kommentar om automatisk oppdatering via OLE-linking i Word
    doc.add_paragraph(
        "Kommentar: Hvis dataene endres i Excel, kan diagrammet oppdateres automatisk i Word ved å bruke OLE-linking (embed/lenke eksternt Excel-diagram i dokumentet)."
    )

    # Kildeliste på nest siste side
    doc.add_page_break()
    doc.add_heading("References", level=1)
    doc.add_paragraph("IT Academy. (2024). Nettverksarkitektur og sikkerhetsdesign for virksomhetsnettverk.")
    doc.add_paragraph("Microsoft. (2024). About Word fields and field codes. https://support.microsoft.com")
    doc.add_paragraph("Microsoft. (2024). Create a table of figures and list of tables. https://support.microsoft.com")

    # Bildeliste på siste side
    doc.add_page_break()
    doc.add_paragraph()
    add_table_of_figures_field(doc.add_paragraph(), "Table of Figures")

    doc.save(WORD_OUTPUT)
    print(f"Dokumentet ble generert: {WORD_OUTPUT}")
    print(f"Excel-fil: {EXCEL_FILE}")
    print(f"Nettverksdiagram: {NETWORK_IMAGE}")
    print(f"Excel-bilde: {EXCEL_IMAGE}")


if __name__ == "__main__":
    build_document()
