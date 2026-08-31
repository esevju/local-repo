#!/usr/bin/env python3
import sys
import os
import traceback

# Write to file since terminal output seems blocked
log_file = r"c:\Users\esevju\OneDrive - Centric IT Academy\VS Code\local-repo\debug_log.txt"

try:
    with open(log_file, 'w') as f:
        f.write("Starting script...\n")
        f.write(f"Python version: {sys.version}\n")
        f.write(f"Working directory: {os.getcwd()}\n")
        f.write("Importing libraries...\n")
        
        from docx import Document
        f.write("✓ Imported Document\n")
        
        from PIL import Image, ImageDraw
        f.write("✓ Imported PIL\n")
        
        from openpyxl import Workbook
        f.write("✓ Imported openpyxl\n")
        
        f.write("\nAll imports successful! Creating document...\n")
        
        # Try creating a simple document
        doc = Document()
        doc.add_heading("Test Document", 0)
        doc.add_paragraph("This is a test")
        
        out_path = r"c:\Users\esevju\OneDrive - Centric IT Academy\VS Code\local-repo\test_output.docx"
        doc.save(out_path)
        f.write(f"✓ Saved test document to: {out_path}\n")
        
except Exception as e:
    with open(log_file, 'a') as f:
        f.write(f"\nERROR: {str(e)}\n")
        f.write(traceback.format_exc())

print("Script completed. Check debug_log.txt for details.")
